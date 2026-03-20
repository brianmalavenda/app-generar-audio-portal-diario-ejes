from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import os
import sys
# If you want to pretty-print the XML (add indentation and line breaks),
# you can use xml.dom.minidom
import xml.dom.minidom

import logging
# Configurar logging para que vaya a stdout (se captura con docker logs)
logging.basicConfig(
    level=logging.DEBUG,  # Nivel de log (DEBUG, INFO, WARNING, ERROR, CRITICAL)
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger(__name__)

class Heading1NotFoundException(Exception):
    """Excepción personalizada para cuando no se encuentra un Heading 1"""
    pass

def extraer_texto_resaltado(input_path, output_path):
    """
    Extrae texto resaltado en amarillo de un documento Word y lo guarda en un nuevo documento.
    Args:
        input_path (str): Ruta al documento Word de entrada (.docx)
        output_dir (str): Directorio donde se guardará el nuevo documento con el texto extraído
    """
    try:
        # Cargar el documento
        doc = Document(input_path)        
        nuevo_doc = Document()
        contador_notas = 0
        nota = []
        
        for paragraph in doc.paragraphs:
            # Iterar a través de todos los "runs" (fragmentos de texto con formato) en el párrafo
            # busca titulos de notas  
            if paragraph.style.name.startswith('Heading 1'):
                if len(nota) == 0:
                    #creo la primera nota
                    nota.append({
                        "id": contador_notas,
                        "titulo": paragraph.text,
                        "cuerpo": [],
                        "estado": "pendiente" # la nota cambiara de estado a completa cuando se acabe el documento o cuando se encuentre otro heading 1
                    })
                else:
                    # Si encontramos otro Heading 1, validamos que la nota tenga cuerpo antes de iniciar una nueva
                    if len(nota[contador_notas]["cuerpo"]) > 0:
                        nota[contador_notas]["estado"] = "completa"
                        # print(f"Nota completa: {nota}")
                        contador_notas += 1
                        # Iniciar una nueva nota
                        nota.append({
                            "id": contador_notas,
                            "titulo": paragraph.text,
                            "cuerpo": [],
                            "estado": "pendiente"
                        })
            for run in paragraph.runs:                              
                # Verificar si el texto está resaltado en amarillo
                if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                    if not paragraph.style.name.startswith('Heading 1'):
                        texto_resaltado = run.text
                        # Añadir el párrafo con su índice
                        nota[contador_notas]['cuerpo'].append({
                            'indice': len(nota[contador_notas]['cuerpo']) + 1,
                            'texto': texto_resaltado
                        })                                          
        
        # Guardar las notas en el nuevo documento
        for nota_a_guardar in nota:
            nuevo_doc.add_heading(nota_a_guardar["titulo"], level=1)
            for parrafo in nota_a_guardar["cuerpo"]:
                nuevo_doc.add_paragraph(f"{parrafo['texto']}")                
        
        # Guardar el documento
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        nuevo_doc.save(output_path)

        logger.info (f"main.py - extraer_texto_resaltado - 01 - ¡Proceso completado! Texto extraído guardado en: {output_path}")
        return output_path  # Devolver la ruta del archivo guardado
        
    except Heading1NotFoundException as e:
        logger.info (f"main.py - extraer_texto_resaltado - 02 - Error: {str(e)}")
        raise
    except Exception as e:
        logger.info (f"main.py - extraer_texto_resaltado - 03 - Error al procesar el documento: {str(e)}")
        raise

def contar_cantidad_de_palabras(input_path):
    """
    Cuenta la cantidad de palabras en un texto dado.
    Args:
        texto (str): El texto en el que se contarán las palabras.
    Returns:
        int: La cantidad de palabras en el texto.
    """
    doc = Document(input_path)        
    nuevo_doc = Document()
    total_palabras = 0
        
    for paragraph in doc.paragraphs:
        texto_parrafo = paragraph.text
        
        # Limpiar caracteres especiales
        texto_parrafo = texto_parrafo.replace(',', '')   # eliminar comas
        texto_parrafo = texto_parrafo.replace('.', '')   # eliminar puntos
        texto_parrafo = texto_parrafo.replace(';', '')   # eliminar punto y coma
        texto_parrafo = texto_parrafo.replace(':', '')   # eliminar dos puntos
        texto_parrafo = texto_parrafo.replace('!', '')   # eliminar exclamaciones
        texto_parrafo = texto_parrafo.replace('?', '')   # eliminar interrogaciones
        
        # Dividir en palabras y contar
        palabras = texto_parrafo.split()
        total_palabras += len(palabras)
    
    return total_palabras

def contar_cantidad_de_caracteres(input_path):
    """
    Cuenta la cantidad de caracteres en un texto dado.
    Args:
        texto (str): El texto en el que se contarán los caracteres.
    Returns:
        int: La cantidad de caracteres en el texto.
    """
    doc = Document(input_path)        
    nuevo_doc = Document()
    total_caracteres = 0

    for paragraph in doc.paragraphs:
        texto_parrafo = paragraph.text
        
        # Limpiar caracteres especiales
        texto_parrafo = texto_parrafo.join(texto_parrafo.split())
        total_caracteres += len(texto_parrafo)

    return total_caracteres

def tamanio_archivo_en_megabytes(ruta_archivo):
    """
    Obtiene el tamaño de un archivo en megabytes.
    Args:
        ruta_archivo (str): La ruta al archivo.
    Returns:
        float: El tamaño del archivo en megabytes.
    """
    if os.path.isfile(ruta_archivo):
        tamaño_bytes = os.path.getsize(ruta_archivo)
        tamaño_megabytes = tamaño_bytes / (1024 * 1024)  # Convertir bytes a megabytes
        return tamaño_megabytes
    else:
        raise FileNotFoundError(f"El archivo XML no existe.")

def convertir_a_formato_ssml(input_path,output_path):
    """
    Convierte el texto del documento de salida a formato SSML.
    Args:
        documento_salida (str): Ruta al documento Word de salida (.docx)
    """
    try:
        doc = Document(input_path)
        ssml_output_string = '<?xml version="1.0"?><speak>'
        
        """
        Encabezado: 
            <?xml version="1.0"?>
            <speak>
        Heading 1 formato SSML:
            <voice name="es-US-Standard-B" gender="MALE">
            <prosody rate="medium" volume="loud">
            <emphasis level="strong">

        Cuerpo formato SSML:
            <voice name="es-US-Standard-A" gender="FEMALE">
            <prosody rate="medium" volume="medium">
        """
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading 1'):
                ssml_output_string += f'<voice name="es-US-Standard-B" gender="MALE"><prosody rate="medium" volume="loud"><emphasis level="strong">{paragraph.text}</emphasis></prosody></voice>\n'
            else:
                ssml_output_string += f'<voice name="es-US-Standard-A" gender="FEMALE"><prosody rate="medium" volume="medium">{paragraph.text}</prosody></voice>'

        ssml_output_string += '</speak>'
    
        # Parse the XML string with minidom
        dom = xml.dom.minidom.parseString(ssml_output_string)
        # ssml_output = dom.toxml()
        # If you want to pretty-print the XML
        ssml_output_pretty_xml = dom.toprettyxml(indent="    ")

        cantidad_palabras = contar_cantidad_de_palabras(ssml_output_pretty_xml)
        cantidad_caracteres = contar_cantidad_de_caracteres(ssml_output_pretty_xml)        
        
        # Guardar el SSML en un archivo
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(ssml_output_pretty_xml)
            
        logger.info (f"main.py - convertir_a_formato_ssml - 01 - ¡SSML generado y guardado en: {output_path}!")
        return [cantidad_palabras, cantidad_caracteres]
    except Exception as e:
        logger.info (f"main.py - convertir_a_formato_ssml - 02 - Error al convertir a SSML: {str(e)}")
    
    # creo que esta parte no tiene sentido
    
    # Obtener el nombre del archivo del cuerpo de la solicitud
    # base_url = "http://localhost:5001/"
    # url = base_url + filename
    # # filename = request.get_json()['filename']
    # file_path = os.path.join(SAVE_FOLDER, filename)

    # # Verificar que el archivo existe
    # if not os.path.exists(file_path):
    #     return jsonify({'error': 'File not found'}), 404

    # response = requests.post(url)

    # # return jsonify({'message': 'Audio generated successfully', 'filename': filename}), 200
    # return send_file(file_path, as_attachment=True)

def leer_archivo_ssml(file_path: str) -> str:
    """Lee archivo SSML/XML y retorna contenido como string"""
    with open(file_path, 'r') as f:
        return f.read()
