# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import os
# If you want to pretty-print the XML (add indentation and line breaks),
# you can use xml.dom.minidom
import xml.dom.minidom

class Heading1NotFoundException(Exception):
    """Excepción personalizada para cuando no se encuentra un Heading 1"""
    pass

def extraer_texto_resaltado(input_path, output_path):
    """
    Extrae texto resaltado en amarillo de un documento Word y lo guarda en un nuevo documento.
    Args:
        input_path (str): Ruta al documento Word de entrada (.docx)
        output_path (str): Ruta donde se guardará el nuevo documento con el texto extraído
    """
    try:
        # Cargar el documento
        doc = Document(input_path)        
        nuevo_doc = Document()
        contador_notas = 0
        nota = []
        
        for paragraph in doc.paragraphs:
            # Iterar a través de todos los "runs" (fragmentos de texto con formato) en el párrafo
            # busca titulos de notas  
            if paragraph.style.name.startswith('Heading 1'):
                if len(nota) == 0:
                    #creo la primera nota
                    nota.append({
                        "id": contador_notas,
                        "titulo": paragraph.text,
                        "cuerpo": [],
                        "estado": "pendiente" # la nota cambiara de estado a completa cuando se acabe el documento o cuando se encuentre otro heading 1
                    })
                else:
                    # Si encontramos otro Heading 1, validamos que la nota tenga cuerpo antes de iniciar una nueva
                    if len(nota[contador_notas]["cuerpo"]) > 0:
                        nota[contador_notas]["estado"] = "completa"
                        # print(f"Nota completa: {nota}")
                        contador_notas += 1
                        # Iniciar una nueva nota
                        nota.append({
                            "id": contador_notas,
                            "titulo": paragraph.text,
                            "cuerpo": [],
                            "estado": "pendiente"
                        })
            for run in paragraph.runs:                              
                # Verificar si el texto está resaltado en amarillo
                if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                    if not paragraph.style.name.startswith('Heading 1'):
                        texto_resaltado = run.text
                        # Añadir el párrafo con su índice
                        nota[contador_notas]['cuerpo'].append({
                            'indice': len(nota[contador_notas]['cuerpo']) + 1,
                            'texto': texto_resaltado
                        })                                          
        # la nota esta completa, antes de continuar con el siguiente párrafo, guardamos la nota en el nuevo documento
        for nota_a_guardar in nota:
            nuevo_doc.add_heading(nota_a_guardar["titulo"], level=1)
            for parrafo in nota_a_guardar["cuerpo"]:
                nuevo_doc.add_paragraph(f"{parrafo['texto']}")                
        nuevo_doc.save(output_path)
        print(f"¡Proceso completado! Texto extraído guardado en: {output_path}")
    except Heading1NotFoundException as e:
        print(f"Error: {str(e)}")
    except Exception as e:
        print(f"Error al procesar el documento: {str(e)}")

def contar_cantidad_de_palabras(texto):
    """
    Cuenta la cantidad de palabras en un texto dado.
    Args:
        texto (str): El texto en el que se contarán las palabras.
    Returns:
        int: La cantidad de palabras en el texto.
    """
    texto = texto.replace(',', '') # eliminar comas para un conteo más preciso
    texto = texto.replace('.', '') # eliminar puntos para un conteo más preciso
    palabras = texto.split()
    return len(palabras)

def contar_cantidad_de_caracteres(texto):
    """
    Cuenta la cantidad de caracteres en un texto dado.
    Args:
        texto (str): El texto en el que se contarán los caracteres.
    Returns:
        int: La cantidad de caracteres en el texto.
    """
    return len(texto)

def tamanio_archivo_en_megabytes(ruta_archivo):
    """
    Obtiene el tamaño de un archivo en megabytes.
    Args:
        ruta_archivo (str): La ruta al archivo.
    Returns:
        float: El tamaño del archivo en megabytes.
    """
    if os.path.isfile(ruta_archivo):
        tamaño_bytes = os.path.getsize(ruta_archivo)
        tamaño_megabytes = tamaño_bytes / (1024 * 1024)  # Convertir bytes a megabytes
        return tamaño_megabytes
    else:
        raise FileNotFoundError(f"El archivo XML no existe.")

def convertir_a_formato_ssml(input_path,output_path):
    """
    Convierte el texto del documento de salida a formato SSML.
    Args:
        documento_salida (str): Ruta al documento Word de salida (.docx)
    """
    try:
        doc = Document(input_path)
        ssml_output_string = '<?xml version="1.0"?><speak>'
        
        """
        Encabezado: 
            <?xml version="1.0"?>
            <speak>
        Heading 1 formato SSML:
            <voice name="es-US-Standard-B" gender="MALE">
            <prosody rate="medium" volume="loud">
            <emphasis level="strong">

        Cuerpo formato SSML:
            <voice name="es-US-Standard-A" gender="FEMALE">
            <prosody rate="medium" volume="medium">
        """
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading 1'):
                ssml_output_string += f'<voice name="es-US-Standard-B" gender="MALE"><prosody rate="medium" volume="loud"><emphasis level="strong">{paragraph.text}</emphasis></prosody></voice>\n'
            else:
                ssml_output_string += f'<voice name="es-US-Standard-A" gender="FEMALE"><prosody rate="medium" volume="medium">{paragraph.text}</prosody></voice>'

        ssml_output_string += '</speak>'
    
        # Parse the XML string with minidom
        dom = xml.dom.minidom.parseString(ssml_output_string)
        # ssml_output = dom.toxml()
        # If you want to pretty-print the XML
        ssml_output_pretty_xml = dom.toprettyxml(indent="    ")

        cantidad_palabras = contar_cantidad_de_palabras(ssml_output_pretty_xml)
        cantidad_caracteres = contar_cantidad_de_caracteres(ssml_output_pretty_xml)        
        
        # Guardar el SSML en un archivo
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(ssml_output_pretty_xml)
            
        print(f"¡SSML generado y guardado en: {output_path}!")
        return [cantidad_palabras, cantidad_caracteres]
    except Exception as e:
        print(f"Error al convertir a SSML: {str(e)}")

if __name__ == "__main__":
    # Configurar rutas de entrada y salida
    # este documento lo ingresa el usuario y el documento de salida tendra el mismo nombre pero con _resaltado
    path_entrada = "/home/brian/Repositorio/app-generar-audio-portal-diario-ejes/diario_pintado/"
    documento_entrada = path_entrada + "test" + ".docx"  # Cambia por la ruta de tu documento
    path_salida = "/home/brian/Repositorio/app-generar-audio-portal-diario-ejes/diario_procesado/"
    documento_salida = path_salida + "test" + "resaltado" + ".docx"
    path_salida = "/home/brian/Repositorio/app-generar-audio-portal-diario-ejes/diario_ssml/"
    xml_salida = path_salida + "test" + "formato_ssml" + ".xml"
    

    # Verificar si el archivo de entrada existe
    if not os.path.exists(documento_entrada):
        print(f"Error: El archivo {documento_entrada} no existe.")
    else:
        # Ejecutar la función principal
        extraer_texto_resaltado(documento_entrada, documento_salida)
        palabras_caracteres = convertir_a_formato_ssml(documento_salida, xml_salida)
        tamanio_megabytes_archivo = tamanio_archivo_en_megabytes(xml_salida)

        print(f"Cantidad de palabras en SSML: {palabras_caracteres[0]}")
        print(f"Cantidad de caracteres en SSML: {palabras_caracteres[1]}")
        print(f"Tamaño del archivo SSML en megabytes: {tamanio_megabytes_archivo}")

