# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import os

class Heading1NotFoundException(Exception):
    """Excepción personalizada para cuando no se encuentra un Heading 1"""
    pass

def extraer_texto_resaltado(input_path, output_path):
    """
    Extrae texto resaltado en amarillo de un documento Word y lo guarda en un nuevo documento.
    Args:
        input_path (str): Ruta al documento Word de entrada (.docx)
        output_path (str): Ruta donde se guardará el nuevo documento con el texto extraído
    """
    try:
        # Cargar el documento
        doc = Document(input_path)        
        nuevo_doc = Document()
        contador_notas = 0
        nota = []
        
        for paragraph in doc.paragraphs:
            # Iterar a través de todos los "runs" (fragmentos de texto con formato) en el párrafo
            # busca tituls de notas  
            print(f" Text: {paragraph.text} - Style: {paragraph.style.name}")
            if paragraph.style.name.startswith('Heading 1'):
                if len(nota) == 0:
                    #creo la primera nota
                    nota.append({
                        "id": contador_notas,
                        "titulo": paragraph.text,
                        "cuerpo": [],
                        "estado": "pendiente" # la nota cambiara de estado a completa cuando se acabe el documento o cuando se encuentre otro heading 1
                    })
                    print(f"Encontrado Heading 1: {nota[contador_notas]["titulo"]}")
                else:
                    # Si encontramos otro Heading 1, validamos que la nota tenga cuerpo antes de iniciar una nueva
                    if len(nota[contador_notas]["cuerpo"]) > 0:
                        nota[contador_notas]["estado"] = "completa"
                        print(f"Nota completa: {nota}")
                        contador_notas += 1
                        # Iniciar una nueva nota
                        nota.append({
                            "id": contador_notas,
                            "titulo": paragraph.text,
                            "cuerpo": [],
                            "estado": "pendiente"
                        })
            for run in paragraph.runs:                              
                # Verificar si el texto está resaltado en amarillo
                if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                    texto_resaltado = run.text
                    # Añadir el párrafo con su índice
                    nota[contador_notas]['cuerpo'].append({
                        'indice': len(nota[contador_notas]['cuerpo']) + 1,
                        'texto': texto_resaltado
                    })                                          
        # la nota esta completa, antes de continuar con el siguiente párrafo, guardamos la nota en el nuevo documento
        print("Guardo la nota en el nuevo documento")
        for nota_a_guardar in nota:
            nuevo_doc.add_heading(nota_a_guardar["titulo"], level=1)
            for parrafo in nota_a_guardar["cuerpo"]:
                nuevo_doc.add_paragraph(f"{parrafo['texto']}")                
        nuevo_doc.save(output_path)
        print(f"¡Proceso completado! Texto extraído guardado en: {output_path}")
    except Heading1NotFoundException as e:
        print(f"Error: {str(e)}")
    except Exception as e:
        print(f"Error al procesar el documento: {str(e)}")

if __name__ == "__main__":
    # Configurar rutas de entrada y salida
    # este documento lo ingresa el usuario y el documento de salida tendra el mismo nombre pero con _resaltado
    path_entrada = "/home/brian/Repositorio/app-generar-audio-portal-diario-ejes/diario_pintado/"
    documento_entrada = path_entrada + "test.docx"  # Cambia por la ruta de tu documento
    path_salida = "/home/brian/Repositorio/app-generar-audio-portal-diario-ejes/diario_procesado/"
    documento_salida = path_salida + "test_resaltado.docx"
    
    # Verificar si el archivo de entrada existe
    if not os.path.exists(documento_entrada):
        print(f"Error: El archivo {documento_entrada} no existe.")
    else:
        # Ejecutar la función principal
        extraer_texto_resaltado(documento_entrada, documento_salida)
