import io
import os
import zipfile
import xml.etree.ElementTree as ET
from docx import Document

# Aca extiendo algunas funciones que hoy en día no son útiles. Vamos a trabajar solo con archivos docx y doc. Los formatos xml, txt, pdf inclusive los dejo de lado por ahora.

def validar_procesar(documento_salida):
    """Procesa archivo detectando automáticamente su tipo"""
    
    print(f"Procesando archivo: {documento_salida}")
    
    
    # Determinar si es SSML (XML) o texto plano
    is_ssml = file_ext in ['.xml', '.ssml']
    
    # Validar SSML si corresponde
    if is_ssml:
        try:
            ET.fromstring(content)
        except ET.ParseError as e:
            raise HTTPException(400, f"XML inválido: {str(e)}")


    # Verificar extensión
    _, ext = os.path.splitext(documento_salida.lower())
    
    try:
        if ext == '.docx':
            # Procesar como DOCX
            return procesar_docx(documento_salida)
            
        elif ext == '.xml':
            # Procesar como XML
            return procesar_xml(documento_salida)

        elif ext == '.pdf':
            # Procesar como PDF
            return procesar_pdf(documento_salida)

        elif ext == '.txt':
            # Procesar como TXT
            return procesar_texto(documento_salida)
            
    except Exception as e:
        print(f"Error procesando archivo: {e}")
        return None

def procesar_docx(ruta_archivo):
    """Procesa archivo DOCX"""
    try:
        doc = Document(ruta_archivo)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
        print("DOCX leído correctamente")
        return text
    except Exception as e:
        print(f"Error con python-docx, intentando método alternativo: {e}")
        return procesar_docx_alternativo(ruta_archivo)

# Definir mejor que es docx alternativo
def procesar_docx_alternativo(ruta_archivo):
    """Método alternativo para leer DOCX"""
    try:
        with zipfile.ZipFile(ruta_archivo, 'r') as docx:
            # Leer el documento principal
            xml_content = docx.read('word/document.xml')
            # Parsear XML y extraer texto
            root = ET.fromstring(xml_content)
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            texts = []
            for elem in root.findall('.//w:t', namespaces):
                if elem.text:
                    texts.append(elem.text)
            return ' '.join(texts)
    except Exception as e:
        print(f"Error método alternativo DOCX: {e}")
        return None

def procesar_xml(ruta_archivo):
    """Procesa archivo XML"""
    try:
        # Leer como binario primero para detectar encoding
        with open(ruta_archivo, 'rb') as f:
            raw_data = f.read()
        
        # Intentar diferentes encodings comunes para XML
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                xml_content = raw_data.decode(encoding)
                # Verificar que sea XML válido
                ET.fromstring(xml_content)
                print(f"XML leído correctamente con encoding: {encoding}")
                return xml_content
            except (UnicodeDecodeError, ET.ParseError):
                continue
        
        print("No se pudo decodificar el XML con los encodings probados")
        return None
        
    except Exception as e:
        print(f"Error procesando XML: {e}")
        return None

def procesar_texto(ruta_archivo):
    """Procesa archivo de texto plano"""
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
    
    for encoding in encodings:
        try:
            with open(ruta_archivo, 'r', encoding=encoding) as f:
                content = f.read()
                print(f"Archivo de texto leído con encoding: {encoding}")
                return content
        except UnicodeDecodeError:
            continue
    
    print("No se pudo leer el archivo con ningún encoding")
    return None

def procesar_pdf(ruta_archivo):
    return "Procesamiento de PDF no implementado"

def leer_docx_completo(file_storage):
    """Lee un archivo DOCX incluyendo texto de tablas"""
    try:
        file_stream = io.BytesIO(file_storage.read())
        doc = Document(file_stream)
        
        texto_completo = []
        
        # Leer párrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                texto_completo.append(paragraph.text)
        
        # Leer tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texto_completo.append(cell.text)
        
        return '\n'.join(texto_completo)
        
    except Exception as e:
        print(f"Error leyendo DOCX: {e}")
        return None
